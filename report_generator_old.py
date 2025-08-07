# Enhanced Report Generator for Deviation Reports
import time
import os
import json
from datetime import datetime
from typing import Dict, Any, Optional, List
from dataclasses import asdict
from pathlib import Path

from deviation_structures import (
    DeviationType, Priority, Status, Department,
    DeviationReport, DeviationInfo, HeaderInfo,
    Investigation, RiskAssessment, RootCauseAnalysis,
    AIGeneratedContent, DeviationDataValidator
)

class EnhancedReportGenerator:
    """
    Generate professional deviation reports in multiple formats (PDF, Word, HTML)
    Enhanced to work with AI-generated content and structured data
    """

    def __init__(self, company_name: str = "CostPlus Drug Company"):
        self.company_name = company_name
        self.template_dir = Path("templates")
        self.output_dir = Path("generated_reports")
        self.output_dir.mkdir(exist_ok=True)

    def generate_html_report(self, deviation_report: DeviationReport, output_path: Optional[str] = None) -> str:
        """Generate an HTML version of the deviation report"""

        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"deviation_report_{timestamp}.html"

        html_content = self._build_html_template(deviation_report)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        print(f"HTML report generated: {output_path}")
        return str(output_path)

    def generate_pdf_report(self, deviation_report: DeviationReport, output_path: Optional[str] = None) -> str:
        """Generate a PDF version of the deviation report"""

        try:
            from weasyprint import HTML, CSS
        except ImportError:
            print("WeasyPrint not installed. Install with: pip install weasyprint")
            print("Generating HTML version instead...")
            return self.generate_html_report(deviation_report, output_path)

        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"deviation_report_{timestamp}.pdf"

        html_content = self._build_html_template(deviation_report)
        css_content = self._get_pdf_css()

        # Convert HTML to PDF
        HTML(string=html_content).write_pdf(
            output_path,
            stylesheets=[CSS(string=css_content)]
        )

        print(f"PDF report generated: {output_path}")
        return str(output_path)

    def generate_word_report(self, deviation_report: DeviationReport, output_path: Optional[str] = None) -> str:
        """Generate a Word document version of the deviation report"""

        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import RGBColor
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            print("Generating HTML version instead...")
            return self.generate_html_report(deviation_report, output_path)

        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"deviation_report_{timestamp}.docx"

        doc = Document()

        # Add header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"{self.company_name} - CONFIDENTIAL"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title
        title = doc.add_heading('Deviation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add report content
        self._add_word_content(doc, deviation_report)

        # Add footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Report Generated: {datetime.now().strftime('%m/%d/%Y %I:%M %p')} | Do not copy and/or distribute without permission."

        doc.save(output_path)
        print(f"Word report generated: {output_path}")
        return str(output_path)

    def generate_all_formats(self, deviation_report: DeviationReport, base_filename: Optional[str] = None) -> Dict[str, str]:
        """Generate the report in all available formats"""

        if not base_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"deviation_report_{timestamp}"

        results = {}

        # Generate HTML
        html_path = self.output_dir / f"{base_filename}.html"
        results['html'] = self.generate_html_report(deviation_report, html_path)

        # Generate PDF
        pdf_path = self.output_dir / f"{base_filename}.pdf"
        results['pdf'] = self.generate_pdf_report(deviation_report, pdf_path)

        # Generate Word
        word_path = self.output_dir / f"{base_filename}.docx"
        results['word'] = self.generate_word_report(deviation_report, word_path)
    

        return results

    def _get_immediate_actions_list(self, report) -> List[str]:
        """Extract immediate actions as a simple list of strings"""
        actions = []

        if hasattr(report, 'immediate_actions') and report.immediate_actions:
            if isinstance(report.immediate_actions, list):
                # Handle simple list
                actions = [str(action) for action in report.immediate_actions if action]
            elif hasattr(report.immediate_actions, 'actions'):
                # Handle ImmediateActions object
                for action in report.immediate_actions.actions:
                    if hasattr(action, 'description'):
                        actions.append(action.description)
                    else:
                        actions.append(str(action))

        return actions

    def _get_ai_generated_content(self, report) -> Dict[str, str]:
        """Extract all AI-generated content for display"""
        ai_content = {}

        if report.ai_generated:
            ai_content['enhanced_description'] = getattr(report.ai_generated, 'enhanced_description', '')
            ai_content['immediate_actions_summary'] = getattr(report.ai_generated, 'immediate_actions_summary', '')  # NEW
            ai_content['ishikawa_analysis'] = getattr(report.ai_generated, 'ishikawa_analysis', '')
            ai_content['risk_assessment'] = getattr(report.ai_generated, 'risk_assessment', '')
            ai_content['regulatory_guidance'] = getattr(report.ai_generated, 'regulatory_guidance', '')
            ai_content['regulatory_capa'] = getattr(report.ai_generated, 'regulatory_capa', '')
            ai_content['root_cause_suggestions'] = getattr(report.ai_generated, 'root_cause_suggestions', [])

        return ai_content

    def _build_html_template(self, report: DeviationReport) -> str:
        """Build the HTML template for the deviation report"""

        # Format dates
        occurrence_date = "Not specified"
        if report.deviation_info.date_of_occurrence:
            occurrence_date = report.deviation_info.date_of_occurrence.strftime("%B %d, %Y")

        generated_date = report.header.report_generated.strftime("%B %d, %Y %I:%M %p")

        # Format departments
        departments = ", ".join([dept.value for dept in report.deviation_info.department]) if report.deviation_info.department else "Not specified"

        # Get AI content
        ai_content = self._get_ai_generated_content(report)

        # Get original user inputs from session data (you'll need to pass this data)
        original_immediate_actions = getattr(report, 'original_immediate_actions', '')
        original_risk_assessment = getattr(report, 'original_risk_assessment', '')

        # Build enhanced deviation description section
        deviation_description_html = f"""
            <div class="section">
                <h2>Deviation Description</h2>
                <div class="description-box">
                    <p><strong>Original Description:</strong></p>
                    <p>{report.deviation_info.description or 'No description provided.'}</p>
                </div>"""

        # Add AI-enhanced description if available
        if ai_content['enhanced_description']:
            deviation_description_html += f"""
                <div class="ai-enhancement">
                    <h3>🤖 AI-Enhanced Technical Analysis</h3>
                    <div class="ai-content">
                        <div class="formatted-content">{self._format_ai_content(ai_content['enhanced_description'])}</div>
                    </div>
                </div>"""

        deviation_description_html += "</div>"

        # Build immediate actions section with original + AI content
        immediate_actions_list = self._get_immediate_actions_list(report)
        # Build immediate actions section with original + AI content ONLY
        immediate_actions_html = f"""
            <div class="section">
                <h2>Immediate Actions</h2>"""

        # Show original user response in grey box
        if original_immediate_actions and original_immediate_actions.lower() not in ['none', 'no', 'n/a', '']:
            immediate_actions_html += f"""
                <div class="original-response-box">
                    <p><strong>Original Response:</strong></p>
                    <p>{original_immediate_actions}</p>
                </div>"""
        else:
            immediate_actions_html += """
                <div class="original-response-box">
                    <p><strong>Original Response:</strong></p>
                    <p>No immediate actions were taken at the time of reporting.</p>
                </div>"""

        # Add AI immediate actions assessment if available (NO documented actions section)
        if ai_content['immediate_actions_summary']:
            immediate_actions_html += f"""
                <div class="ai-enhancement">
                    <h3>🤖 AI Assessment of Immediate Actions</h3>
                    <div class="ai-content">
                        <div class="formatted-content">{self._format_ai_content(ai_content['immediate_actions_summary'])}</div>
                    </div>
                </div>"""

        immediate_actions_html += "</div>"

        # Build risk assessment section with original + AI content
        risk_assessment_html = f"""
            <div class="section">
                <h2>Risk Assessment Information</h2>"""

        # Show original user risk assessment in grey box
        if original_risk_assessment and original_risk_assessment.strip():
            risk_assessment_html += f"""
                <div class="original-response-box">
                    <p><strong>Initial Risk Assessment (User Input):</strong></p>
                    <p>{original_risk_assessment}</p>
                </div>"""
        else:
            risk_assessment_html += """
                <div class="original-response-box">
                    <p><strong>Initial Risk Assessment (User Input):</strong></p>
                    <p>No initial risk assessment provided.</p>
                </div>"""

        # Add AI risk assessment if available (NO formal risk assessment section)
        if ai_content['risk_assessment']:
            risk_assessment_html += f"""
                <div class="ai-enhancement">
                    <h3>🤖 AI Risk Assessment</h3>
                    <div class="ai-content">
                        <div class="formatted-content">{self._format_ai_content(ai_content['risk_assessment'])}</div>
                    </div>
                </div>"""

        risk_assessment_html += "</div>"

        # Build additional AI analysis sections (kept separate)
        additional_ai_sections = ""

        if ai_content['ishikawa_analysis']:
            additional_ai_sections += f"""
            <div class="ai-section">
                <h2>🔍 Ishikawa (Fishbone) Root Cause Analysis</h2>
                <div class="ai-content">
                    <div class="formatted-content">{self._format_ai_content(ai_content['ishikawa_analysis'])}</div>
                </div>
            </div>
            """

        if ai_content['root_cause_suggestions']:
            additional_ai_sections += f"""
            <div class="ai-section">
                <h2>🎯 Traditional Root Cause Assessment</h2>
                <div class="ai-content">
                    {"".join([f"<div class='formatted-content'>{self._format_ai_content(suggestion)}</div>" for suggestion in ai_content['root_cause_suggestions']])}
                </div>
            </div>
            """

        if ai_content['regulatory_guidance']:
            additional_ai_sections += f"""
            <div class="ai-section">
                <h2>📋 Regulatory Guidance</h2>
                <div class="ai-content">
                    <div class="formatted-content">{self._format_ai_content(ai_content['regulatory_guidance'])}</div>
                </div>
            </div>
            """

        if ai_content['regulatory_capa']:
            additional_ai_sections += f"""
            <div class="ai-section">
                <h2>⚖️ Regulatory-Compliant CAPA</h2>
                <div class="ai-content">
                    <div class="formatted-content">{self._format_ai_content(ai_content['regulatory_capa'])}</div>
                </div>
            </div>
            """

        # Root cause analysis section (if manually filled)
        root_cause_html = ""
        if report.root_cause_analysis and (report.root_cause_analysis.primary_cause or report.root_cause_analysis.detailed_analysis):
            root_cause_html = f"""
            <div class="section">
                <h2>Root Cause Analysis</h2>
                <table class="info-table">
                    <tr>
                        <td class="label"><strong>Primary Cause:</strong></td>
                        <td>{report.root_cause_analysis.primary_cause.value if report.root_cause_analysis.primary_cause else 'Under investigation'}</td>
                    </tr>
                    <tr>
                        <td class="label"><strong>Human Error Type:</strong></td>
                        <td>{report.root_cause_analysis.human_error_category.value if report.root_cause_analysis.human_error_category else 'Not applicable'}</td>
                    </tr>
                </table>
                <div class="description-box">
                    <p><strong>Detailed Analysis:</strong></p>
                    <p>{report.root_cause_analysis.detailed_analysis or 'Analysis in progress.'}</p>
                </div>
            </div>
            """

        html_template = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Deviation Report - {report.header.record_id}</title>
            <style>
                {self._get_html_css()}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="company-logo">
                    <h1>{self.company_name}</h1>
                </div>
                <div class="confidential">
                    <p>Do not copy and/or distribute without permission.<br>
                    <strong>CONFIDENTIAL</strong></p>
                </div>
            </div>

            <div class="report-title">
                <h1>Deviation Report</h1>
                <div class="report-info">
                    <p><strong>Record ID:</strong> {report.header.record_id}</p>
                    <p><strong>Page:</strong> 1 of 1</p>
                </div>
            </div>

            <div class="section">
                <h2>Deviation Information</h2>
                <table class="info-table">
                    <tr>
                        <td class="label"><strong>Title:</strong></td>
                        <td>{report.deviation_info.title}</td>
                        <td class="label"><strong>Priority:</strong></td>
                        <td class="priority-{report.deviation_info.priority.value.lower()}">{report.deviation_info.priority.value}</td>
                    </tr>
                    <tr>
                        <td class="label"><strong>Status:</strong></td>
                        <td>{report.deviation_info.status.value}</td>
                        <td class="label"><strong>Deviation Type:</strong></td>
                        <td>{report.deviation_info.deviation_type.value if report.deviation_info.deviation_type else 'Not specified'}</td>
                    </tr>
                    <tr>
                        <td class="label"><strong>Date of Occurrence:</strong></td>
                        <td>{occurrence_date}</td>
                        <td class="label"><strong>Department(s):</strong></td>
                        <td>{departments}</td>
                    </tr>
                    <tr>
                        <td class="label"><strong>Batch/Lot Number:</strong></td>
                        <td>{report.deviation_info.batch_lot_number or 'Not applicable'}</td>
                        <td class="label"><strong>Quantity Impacted:</strong></td>
                        <td>{report.deviation_info.quantity_impacted or 'Not specified'}</td>
                    </tr>
                    <tr>
                        <td class="label"><strong>Planned Deviation:</strong></td>
                        <td colspan="3">{'Yes' if report.deviation_info.is_planned_deviation else 'No'}</td>
                    </tr>
                </table>
            </div>

            {deviation_description_html}

            {immediate_actions_html}

            {risk_assessment_html}

            {root_cause_html}

            {additional_ai_sections}

            <div class="section">
                <h2>Investigation Information</h2>
                <table class="info-table">
                    <tr>
                        <td class="label"><strong>Investigation Assignee:</strong></td>
                        <td>{report.investigation.assignee or 'Not assigned'}</td>
                    </tr>
                    <tr>
                        <td class="label"><strong>Initiator:</strong></td>
                        <td>{report.header.initiator or 'Not specified'}</td>
                    </tr>
                </table>

                <div class="investigation-box">
                    <p><strong>Investigation Results:</strong></p>
                    <p>{report.investigation.findings or 'Investigation pending.'}</p>

                    <p><strong>Final Conclusion:</strong></p>
                    <p>{report.investigation.conclusion or 'Investigation in progress.'}</p>
                </div>
            </div>

            <div class="footer">
                <p>Report Generated: {generated_date}</p>
                <p>Do not copy and/or distribute without permission. CONFIDENTIAL</p>
            </div>
        </body>
        </html>
        """
        return html_template

    def _format_ai_content(self, content: str) -> str:
        """Format AI-generated content for HTML display"""
        if not content:
            return ""

        # Remove any markdown headers (##, ###) but keep the text
        content = content.replace('### ', '').replace('## ', '').replace('# ', '')
        
        # Convert numbered lists with proper formatting
        lines = content.split('\n')
        formatted_lines = []
        in_list = False
        
        for line in lines:
            stripped = line.strip()
            
            # Handle numbered items (1., 2., etc.)
            import re
            if re.match(r'^\d+\.', stripped):
                if not in_list:
                    formatted_lines.append('<ol>')
                    in_list = True
                # Extract the content after the number
                list_content = re.sub(r'^\d+\.\s*', '', stripped)
                # Remove any ** markdown bold markers
                list_content = list_content.replace('**', '')
                formatted_lines.append(f'<li>{list_content}</li>')
            # Handle bullet points
            elif stripped.startswith('- ') or stripped.startswith('• '):
                if in_list and formatted_lines[-1] == '</ol>':
                    formatted_lines.pop()  # Remove the </ol>
                    formatted_lines.append('<ul>')
                elif not in_list:
                    formatted_lines.append('<ul>')
                    in_list = True
                bullet_content = stripped[2:].strip().replace('**', '')
                formatted_lines.append(f'<li>{bullet_content}</li>')
            else:
                if in_list:
                    if formatted_lines[-1].startswith('<li>'):
                        formatted_lines.append('</ol>' if '<ol>' in ''.join(formatted_lines) else '</ul>')
                    in_list = False
                if stripped:
                    # Remove any ** markdown bold markers and add as regular paragraph
                    clean_text = stripped.replace('**', '')
                    formatted_lines.append(f'<p style="color: #333; font-weight: normal;">{clean_text}</p>')
        
        if in_list:
            formatted_lines.append('</ol>' if '<ol>' in ''.join(formatted_lines) else '</ul>')
        
        return '\n'.join(formatted_lines)

    def _get_html_css(self) -> str:
        """Get CSS styles for HTML report"""
        return """
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }

            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 210mm;
                margin: 0 auto;
                padding: 20px;
                background: white;
            }

            .header {
                display: flex;
                justify-content: space-between;
                align-items: center;
                border-bottom: 3px solid #0066cc;
                padding-bottom: 15px;
                margin-bottom: 30px;
            }

            .company-logo h1 {
                color: #0066cc;
                font-size: 24px;
                font-weight: bold;
            }

            .confidential {
                text-align: right;
                font-size: 12px;
                color: #666;
            }

            .report-title {
                text-align: center;
                margin-bottom: 30px;
            }

            .report-title h1 {
                font-size: 28px;
                color: #0066cc;
                margin-bottom: 10px;
            }

            .report-info {
                display: flex;
                justify-content: center;
                gap: 40px;
                font-size: 14px;
            }

            .section {
                margin-bottom: 25px;
                border: 1px solid #ddd;
                border-radius: 5px;
                overflow: hidden;
            }

            .section h2 {
                background: #f8f9fa;
                color: #333;
                padding: 12px 15px;
                margin: 0;
                border-bottom: 1px solid #ddd;
                font-size: 18px;
            }

            /* Original user response boxes (grey) */
            .original-response-box {
                padding: 15px;
                background: #f5f5f5;
                border: 1px solid #ddd;
                margin: 10px;
                border-radius: 5px;
                border-left: 4px solid #6c757d;
            }

            .original-response-box p {
                margin: 5px 0;
            }

            .original-response-box strong {
                color: #495057;
            }

            /* AI Enhancement sections within main sections (green) */
            .ai-enhancement {
                margin: 15px;
                border: 2px solid #28a745;
                border-radius: 8px;
                overflow: hidden;
                background: #f8fff8;
            }

            .ai-enhancement h3 {
                background: #28a745;
                color: white;
                padding: 8px 12px;
                margin: 0;
                font-size: 14px;
                font-weight: bold;
            }

            /* Standalone AI analysis sections */
            .ai-section {
                margin-bottom: 20px;
                border: 2px solid #4CAF50;
                border-radius: 8px;
                overflow: hidden;
                background: #f9fff9;
            }

            .ai-section h2 {
                background: #4CAF50;
                color: white;
                padding: 12px 15px;
                margin: 0;
                font-size: 16px;
                border-bottom: none;
            }

            .ai-content {
                padding: 12px;
                background: #f8fff8;
            }

            .ai-content p {
                color: #333 !important;
                font-weight: normal !important;
            }

            .ai-content li {
                color: #333 !important;
                font-weight: normal !important;
            }

            .ai-content ol, .ai-content ul {
                color: #333 !important;
                font-weight: normal !important;
            }

            /* Ensure formatted content is not bold */
            .formatted-content * {
                font-weight: normal !important;
            }

            .info-table {
                width: 100%;
                border-collapse: collapse;
            }

            .info-table td {
                padding: 12px 15px;
                border-bottom: 1px solid #eee;
                vertical-align: top;
            }

            .info-table td.label {
                background: #f8f9fa;
                font-weight: bold;
                width: 20%;
                color: #555;
            }

            .priority-critical {
                background: #dc3545;
                color: white;
                padding: 4px 8px;
                border-radius: 4px;
                font-weight: bold;
            }

            .priority-major {
                background: #fd7e14;
                color: white;
                padding: 4px 8px;
                border-radius: 4px;
                font-weight: bold;
            }

            .priority-minor {
                background: #28a745;
                color: white;
                padding: 4px 8px;
                border-radius: 4px;
                font-weight: bold;
            }

            .description-box, .actions-box, .risk-box, .investigation-box {
                padding: 15px;
                background: #fafafa;
                border: 1px solid #eee;
                margin: 10px;
            }

            .actions-box ul {
                margin-left: 20px;
            }

            .actions-box li {
                margin-bottom: 8px;
            }

            .footer {
                margin-top: 40px;
                padding-top: 20px;
                border-top: 2px solid #0066cc;
                text-align: center;
                font-size: 12px;
                color: #666;
            }

            .formatted-content h4 {
                color: #0066cc;
                font-size: 16px;
                margin: 15px 0 8px 0;
                border-bottom: 1px solid #ddd;
                padding-bottom: 5px;
            }

            .formatted-content h5 {
                color: #555;
                font-size: 14px;
                margin: 12px 0 6px 0;
                font-weight: bold;
            }

            .formatted-content ul {
                margin: 10px 0 10px 20px;
            }

            .formatted-content li {
                margin-bottom: 5px;
                line-height: 1.5;
            }

            .formatted-content p {
                margin: 8px 0;
                line-height: 1.6;
            }

            @media print {
                body {
                    max-width: none;
                    margin: 0;
                    padding: 15px;
                }

                .footer {
                    position: fixed;
                    bottom: 0;
                    left: 0;
                    right: 0;
                    background: white;
                }
            }
        """


    def _get_pdf_css(self) -> str:
        """Get CSS styles optimized for PDF generation"""
        return self._get_html_css() + """
            @page {
                size: A4;
                margin: 2cm;
            }

            body {
                font-size: 11pt;
            }

            /* Prevent sections and content blocks from breaking across pages */
            .section {
                page-break-inside: avoid;
                break-inside: avoid;
            }

            /* Prevent AI sections from breaking */
            .ai-section {
                page-break-inside: avoid;
                break-inside: avoid;
            }

            .ai-enhancement {
                page-break-inside: avoid;
                break-inside: avoid;
            }

            /* Prevent content boxes from breaking */
            .description-box,
            .actions-box,
            .risk-box,
            .investigation-box,
            .original-response-box {
                page-break-inside: avoid;
                break-inside: avoid;
            }

            /* Prevent paragraphs in formatted content from breaking */
            .formatted-content p {
                page-break-inside: avoid;
                break-inside: avoid;
                orphans: 3;
                widows: 3;
            }

            /* Prevent AI content blocks from breaking */
            .ai-content {
                page-break-inside: avoid;
                break-inside: avoid;
            }

            /* Keep tables together */
            .info-table {
                page-break-inside: avoid;
                break-inside: avoid;
            }

            /* Keep headings with their content */
            h1, h2, h3, h4, h5, h6 {
                page-break-after: avoid;
                break-after: avoid;
                orphans: 3;
                widows: 3;
            }

            /* Ensure text blocks stay together */
            p {
                orphans: 3;
                widows: 3;
            }

            /* Keep list items together */
            li {
                page-break-inside: avoid;
                break-inside: avoid;
            }

            /* Keep the entire regulatory CAPA section together if possible */
            .ai-section:has(h2:contains("Regulatory-Compliant CAPA")) {
                page-break-before: auto;
                page-break-after: auto;
            }
        """

    def _add_word_content(self, doc, report: DeviationReport):
        """Add content to Word document"""
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Header information table
        doc.add_heading('Report Information', level=1)
        table = doc.add_table(rows=3, cols=4)
        table.style = 'Table Grid'

        # Header row
        cells = table.rows[0].cells
        cells[0].text = 'Record ID'
        cells[1].text = report.header.record_id
        cells[2].text = 'Initiator'
        cells[3].text = report.header.initiator

        cells = table.rows[1].cells
        cells[0].text = 'Report Generated'
        cells[1].text = report.header.report_generated.strftime("%m/%d/%Y %I:%M %p")
        cells[2].text = 'Page'
        cells[3].text = '1 of 1'

        # Deviation Information
        doc.add_heading('Deviation Information', level=1)

        dev_table = doc.add_table(rows=6, cols=4)
        dev_table.style = 'Table Grid'

        # Populate deviation table
        rows_data = [
            ['Title', report.deviation_info.title, 'Priority', report.deviation_info.priority.value],
            ['Status', report.deviation_info.status.value, 'Type', report.deviation_info.deviation_type.value if report.deviation_info.deviation_type else 'Not specified'],
            ['Date of Occurrence', report.deviation_info.date_of_occurrence.strftime("%m/%d/%Y") if report.deviation_info.date_of_occurrence else 'Not specified', 'Department', ", ".join([d.value for d in report.deviation_info.department])],
            ['Batch/Lot Number', report.deviation_info.batch_lot_number or 'Not applicable', 'Quantity Impacted', report.deviation_info.quantity_impacted or 'Not specified'],
            ['Planned Deviation', 'Yes' if report.deviation_info.is_planned_deviation else 'No', '', '']
        ]

        for i, (col1, val1, col2, val2) in enumerate(rows_data):
            cells = dev_table.rows[i].cells
            cells[0].text = col1
            cells[1].text = val1
            cells[2].text = col2
            cells[3].text = val2

        # Description with AI enhancement
        doc.add_heading('Deviation Description', level=1)
        doc.add_paragraph(report.deviation_info.description or 'No description provided.')

        ai_content = self._get_ai_generated_content(report)
        if ai_content['enhanced_description']:
            doc.add_heading('AI-Enhanced Technical Analysis', level=2)
            doc.add_paragraph(ai_content['enhanced_description'])

        # Immediate Actions with original response + AI assessment ONLY
        doc.add_heading('Immediate Actions', level=1)

        # Show original user response
        original_immediate_actions = getattr(report, 'original_immediate_actions', '')
        if original_immediate_actions and original_immediate_actions.lower() not in ['none', 'no', 'n/a', '']:
            doc.add_paragraph('Original Response:', style='Heading 2')
            doc.add_paragraph(original_immediate_actions)
        else:
            doc.add_paragraph('Original Response:', style='Heading 2')
            doc.add_paragraph('No immediate actions were taken at the time of reporting.')

        # Show AI assessment (NO documented actions section)
        if ai_content['immediate_actions_summary']:
            doc.add_heading('AI Assessment of Immediate Actions', level=2)
            doc.add_paragraph(ai_content['immediate_actions_summary'])

        # Risk Assessment with original + AI content ONLY
        doc.add_heading('Risk Assessment Information', level=1)

        # Show original user risk assessment
        original_risk_assessment = getattr(report, 'original_risk_assessment', '')
        if original_risk_assessment and original_risk_assessment.strip():
            doc.add_paragraph('Initial Risk Assessment (User Input):', style='Heading 2')
            doc.add_paragraph(original_risk_assessment)
        else:
            doc.add_paragraph('Initial Risk Assessment (User Input):', style='Heading 2')
            doc.add_paragraph('No initial risk assessment provided.')

        # Show AI risk assessment (NO formal risk assessment section)
        if ai_content['risk_assessment']:
            doc.add_heading('AI Risk Assessment', level=2)
            doc.add_paragraph(ai_content['risk_assessment'])

        # Additional AI Analysis Sections
        if ai_content['ishikawa_analysis']:
            doc.add_heading('Ishikawa (Fishbone) Root Cause Analysis', level=1)
            doc.add_paragraph(ai_content['ishikawa_analysis'])

        if ai_content['root_cause_suggestions']:
            doc.add_heading('Traditional Root Cause Assessment', level=1)
            for suggestion in ai_content['root_cause_suggestions']:
                doc.add_paragraph(suggestion)

        if ai_content['regulatory_guidance']:
            doc.add_heading('Regulatory Guidance', level=1)
            doc.add_paragraph(ai_content['regulatory_guidance'])

        if ai_content['regulatory_capa']:
            doc.add_heading('Regulatory-Compliant CAPA', level=1)
            doc.add_paragraph(ai_content['regulatory_capa'])

        # Investigation
        doc.add_heading('Investigation Information', level=1)
        doc.add_paragraph(f"Investigation Assignee: {report.investigation.assignee or 'Not assigned'}")
        doc.add_paragraph(f"Findings: {report.investigation.findings or 'Investigation pending.'}")
        doc.add_paragraph(f"Conclusion: {report.investigation.conclusion or 'Investigation in progress.'}")



# Integration function to generate reports from agent sessions
def generate_report_from_agent_session(agent: 'DeviationInterviewAgent', session_id: str, output_format: str = "all") -> Dict[str, str]:
    """
    Generate a deviation report from an AI agent interview session

    Args:
        agent: The DeviationInterviewAgent instance
        session_id: The interview session ID
        output_format: "html", "pdf", "word", or "all"

    Returns:
        Dictionary with file paths of generated reports
    """

    # Get the deviation report from the session
    report = agent.get_session_report(session_id)
    if not report:
        raise ValueError(f"Session {session_id} not found or report not generated")

    # Enhance the report with missing information if needed
    _enhance_report_for_generation(report, agent, session_id)

    # Initialize the report generator
    generator = EnhancedReportGenerator()

    # Generate reports based on requested format
    if output_format.lower() == "all":
        return generator.generate_all_formats(report)
    elif output_format.lower() == "html":
        path = generator.generate_html_report(report)
        return {"html": path}
    elif output_format.lower() == "pdf":
        path = generator.generate_pdf_report(report)
        return {"pdf": path}
    elif output_format.lower() == "word":
        path = generator.generate_word_report(report)
        return {"word": path}
    else:
        raise ValueError("output_format must be 'html', 'pdf', 'word', or 'all'")


def _enhance_report_for_generation(report: DeviationReport, agent: 'DeviationInterviewAgent', session_id: str):
    """Enhance the report with any missing information before generation"""

    # Get session data to fill in any gaps
    session_data = agent.export_session_data(session_id)
    collected_info = session_data.get("collected_info", {})
    structured_data = session_data.get("structured_data", {})

    # Set default values if missing
    if not report.header.record_id:
        report.header.record_id = f"DEV-{session_id}"

    if not report.header.initiator:
        report.header.initiator = "AI Interview System"

    # Set default status and priority if not specified
    if report.deviation_info.status == Status.OPEN:
        report.deviation_info.status = Status.IN_PROGRESS

    # Use enhanced description if available
    if not report.deviation_info.description and "enhanced_description" in collected_info:
        report.deviation_info.description = collected_info["enhanced_description"]

    # Store original user inputs for display in report (NEW)
    report.original_immediate_actions = structured_data.get("immediate_actions", "")
    report.original_risk_assessment = structured_data.get("initial_risk", "")

    # Add immediate actions summary to AI generated content
    if "immediate_actions_summary" in collected_info:
        report.ai_generated.immediate_actions_summary = collected_info["immediate_actions_summary"]

    # Add AI-generated content to risk assessment if empty
    if not report.risk_assessment.conclusion:
        if hasattr(report.ai_generated, 'risk_assessment') and report.ai_generated.risk_assessment:
            report.risk_assessment.conclusion = "Based on comprehensive AI analysis including user's initial assessment and regulatory guidance."
        elif report.ai_generated.root_cause_suggestions:
            report.risk_assessment.conclusion = "Based on AI analysis, see root cause assessment for detailed evaluation."


