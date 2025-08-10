# simplified_report_generator.py - Clean and Simple Report Generator

import os
import json
import html
from datetime import datetime
from typing import Dict, Any, Optional, List
from pathlib import Path

from deviation_structures import (
    DeviationType, Priority, Status, Department,
    DeviationReport, DeviationInfo, HeaderInfo,
    Investigation, RiskAssessment, RootCauseAnalysis,
    AIGeneratedContent, DeviationDataValidator
)

class SimplifiedReportGenerator:
    """
    Generate clean and simple deviation reports in multiple formats
    Focus on User Input vs AI-Generated content separation
    """

    def __init__(self, company_name: str = "CostPlus Drug Company"):
        self.company_name = company_name
        self.output_dir = Path("generated_reports")
        self.output_dir.mkdir(exist_ok=True)

    def generate_html_report(self, deviation_report: DeviationReport, output_path: Optional[str] = None) -> str:
        """Generate an HTML version of the deviation report"""
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"deviation_report_{timestamp}.html"

        html_content = self._build_html_template(deviation_report)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        print(f"HTML report generated: {output_path}")
        return str(output_path)

    def generate_pdf_report(self, deviation_report: DeviationReport, output_path: Optional[str] = None) -> str:
        """Generate a PDF version of the deviation report"""
        try:
            from weasyprint import HTML, CSS
        except ImportError:
            print("WeasyPrint not installed. Install with: pip install weasyprint")
            print("Generating HTML version instead...")
            return self.generate_html_report(deviation_report, output_path)

        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"deviation_report_{timestamp}.pdf"

        html_content = self._build_html_template(deviation_report)
        css_content = self._get_pdf_css()

        HTML(string=html_content).write_pdf(
            output_path,
            stylesheets=[CSS(string=css_content)]
        )

        print(f"PDF report generated: {output_path}")
        return str(output_path)

    def generate_word_report(self, deviation_report: DeviationReport, output_path: Optional[str] = None) -> str:
        """Generate a Word document version of the deviation report"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import RGBColor
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            print("Generating HTML version instead...")
            return self.generate_html_report(deviation_report, output_path)

        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"deviation_report_{timestamp}.docx"

        doc = Document()
        self._add_word_content(doc, deviation_report)
        doc.save(output_path)
        
        print(f"Word report generated: {output_path}")
        return str(output_path)

    def generate_all_formats(self, deviation_report: DeviationReport, base_filename: Optional[str] = None) -> Dict[str, str]:
        """Generate the report in all available formats"""
        if not base_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"deviation_report_{timestamp}"

        results = {}

        # Generate HTML
        html_path = self.output_dir / f"{base_filename}.html"
        results['html'] = self.generate_html_report(deviation_report, html_path)

        # Generate PDF
        pdf_path = self.output_dir / f"{base_filename}.pdf"
        results['pdf'] = self.generate_pdf_report(deviation_report, pdf_path)

        # Generate Word
        word_path = self.output_dir / f"{base_filename}.docx"
        results['word'] = self.generate_word_report(deviation_report, word_path)

        return results

    def _extract_ai_sections(self, ai_content: str) -> Dict[str, str]:
        """Extract individual sections from AI-generated content"""
        sections = {}
        
        if not ai_content:
            return sections
        
        # Define section patterns to look for
        section_patterns = [
            ("Deviation Summary", ["**Deviation Summary**", "**1. Deviation Summary**", "1. **Deviation Summary**"]),
            ("Event Timeline", ["**Event Timeline**", "**2. Event Timeline**", "2. **Event Timeline**"]),
            ("Root Cause Analysis", ["**Root Cause Analysis**", "**3. Root Cause Analysis**", "3. **Root Cause Analysis**"]),
            ("Impact Assessment", ["**Impact Assessment**", "**4. Impact Assessment**", "4. **Impact Assessment**"]),
            ("CAPA Plan", ["**Corrective and Preventive Action (CAPA) Plan**", "**5. Corrective and Preventive Action (CAPA) Plan**", "5. **Corrective and Preventive Action (CAPA) Plan**", "**CAPA Plan**"])
        ]
        
        lines = ai_content.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            
            # Check if this line starts a new section
            found_section = None
            for section_name, patterns in section_patterns:
                for pattern in patterns:
                    if pattern in line:
                        found_section = section_name
                        break
                if found_section:
                    break
            
            if found_section:
                # Save previous section if exists
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                # Start new section
                current_section = found_section
                current_content = []
            elif current_section and line:
                # Add content to current section (skip the header line)
                if not any(pattern in line for _, patterns in section_patterns for pattern in patterns):
                    current_content.append(line)
        
        # Save the last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections

    def _build_html_template(self, report: DeviationReport) -> str:
        """Build the HTML template for the deviation report"""
        
        # Format dates
        occurrence_date = "Not specified"
        if report.deviation_info.date_of_occurrence:
            occurrence_date = report.deviation_info.date_of_occurrence.strftime("%B %d, %Y")

        generated_date = report.header.report_generated.strftime("%B %d, %Y %I:%M %p")

        # Format departments
        departments = ", ".join([dept.value for dept in report.deviation_info.department]) if report.deviation_info.department else "Not specified"

        # Extract AI sections
        ai_content = getattr(report.ai_generated, 'enhanced_description', '')
        ai_sections = self._extract_ai_sections(ai_content)

        # Build AI sections HTML with blue headers
        ai_sections_html = ""
        
        section_order = ["Deviation Summary", "Event Timeline", "Root Cause Analysis", "Impact Assessment", "CAPA Plan"]
        section_icons = {
            "Deviation Summary": "📋",
            "Event Timeline": "⏰", 
            "Root Cause Analysis": "🔍",
            "Impact Assessment": "⚠️",
            "CAPA Plan": "🔧"
        }
        
        for section_name in section_order:
            if section_name in ai_sections and ai_sections[section_name]:
                icon = section_icons.get(section_name, "🤖")
                ai_sections_html += f"""
                <div class="ai-section">
                    <h2 class="ai-header">{icon} {section_name}</h2>
                    <div class="ai-content">
                        {self._format_ai_content(ai_sections[section_name])}
                    </div>
                </div>
                """

        # Get FULL user description - not truncated
        full_user_description = report.deviation_info.description or 'No description provided.'
        
        # Format the description to preserve line breaks
        formatted_user_description = html.escape(full_user_description)
        formatted_user_description = formatted_user_description.replace('\n', '<br>')
        formatted_user_description = formatted_user_description.replace('\r\n', '<br>')
        formatted_user_description = formatted_user_description.replace('\r', '<br>')


        html_template = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Deviation Report - {report.header.record_id}</title>
            <style>
                {self._get_html_css()}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="company-logo">
                    <h1>{self.company_name}</h1>
                </div>
                <div class="confidential">
                    <p>Do not copy and/or distribute without permission.<br>
                    <strong>CONFIDENTIAL</strong></p>
                </div>
            </div>

            <div class="report-title">
                <h1>Deviation Report</h1>
                <div class="report-info">
                    <p><strong>Record ID:</strong> {report.header.record_id}</p>
                    <p><strong>Generated:</strong> {generated_date}</p>
                </div>
            </div>

            <div class="section">
                <h2>📊 Report Information</h2>
                <table class="info-table">
                    <tr>
                        <td class="label"><strong>Title:</strong></td>
                        <td>{report.deviation_info.title}</td>
                        <td class="label"><strong>Priority:</strong></td>
                        <td class="priority-{report.deviation_info.priority.value.lower()}">{report.deviation_info.priority.value}</td>
                    </tr>
                    <tr>
                        <td class="label"><strong>Status:</strong></td>
                        <td>{report.deviation_info.status.value}</td>
                        <td class="label"><strong>Deviation Type:</strong></td>
                        <td>{report.deviation_info.deviation_type.value if report.deviation_info.deviation_type else 'Not specified'}</td>
                    </tr>
                    <tr>
                        <td class="label"><strong>Date of Occurrence:</strong></td>
                        <td>{occurrence_date}</td>
                        <td class="label"><strong>Department(s):</strong></td>
                        <td>{departments}</td>
                    </tr>
                    <tr>
                        <td class="label"><strong>Batch/Lot Number:</strong></td>
                        <td>{report.deviation_info.batch_lot_number or 'Not applicable'}</td>
                        <td class="label"><strong>Quantity Impacted:</strong></td>
                        <td>{report.deviation_info.quantity_impacted or 'Not specified'}</td>
                    </tr>
                    <tr>
                        <td class="label"><strong>Planned Deviation:</strong></td>
                        <td>{'Yes' if report.deviation_info.is_planned_deviation else 'No'}</td>
                        <td class="label"><strong>Initiator:</strong></td>
                        <td>{report.header.initiator or 'Not specified'}</td>
                    </tr>
                </table>
            </div>

            <div class="section user-input-section">
                <h2 class="user-input-header">👤 User Input - Original Deviation Description</h2>
                <div class="user-input-box">
                    <div class="full-description">
                        {formatted_user_description}
                    </div>
                </div>
            </div>

            <div class="ai-generated-container">
                <h2 class="ai-main-header">🤖 AI-Generated Analysis</h2>
                {ai_sections_html}
            </div>

            <div class="footer">
                <p>Report Generated: {generated_date} | {self.company_name}</p>
                <p>CONFIDENTIAL - Do not copy and/or distribute without permission</p>
            </div>
        </body>
        </html>
        """
        return html_template

    def _format_ai_content(self, content: str) -> str:
        """Format AI-generated content for HTML display"""
        if not content:
            return ""

        # Clean up markdown formatting
        content = content.replace('**', '').replace('##', '').replace('#', '')
        
        # Convert to HTML with proper formatting
        lines = content.split('\n')
        formatted_lines = []
        in_list = False
        
        for line in lines:
            stripped = line.strip()
            
            # Handle numbered items (1., 2., etc.)
            import re
            if re.match(r'^\d+\.', stripped):
                if not in_list:
                    formatted_lines.append('<ol>')
                    in_list = True
                list_content = re.sub(r'^\d+\.\s*', '', stripped)
                formatted_lines.append(f'<li>{list_content}</li>')
            # Handle bullet points
            elif stripped.startswith('- ') or stripped.startswith('• '):
                if in_list and formatted_lines and not formatted_lines[-1].startswith('<li>'):
                    formatted_lines.append('</ol>')
                    formatted_lines.append('<ul>')
                elif not in_list:
                    formatted_lines.append('<ul>')
                    in_list = True
                bullet_content = stripped[2:].strip()
                formatted_lines.append(f'<li>{bullet_content}</li>')
            else:
                if in_list:
                    if '<ol>' in ''.join(formatted_lines[-10:]):
                        formatted_lines.append('</ol>')
                    else:
                        formatted_lines.append('</ul>')
                    in_list = False
                if stripped:
                    formatted_lines.append(f'<p>{stripped}</p>')
        
        if in_list:
            if '<ol>' in ''.join(formatted_lines[-10:]):
                formatted_lines.append('</ol>')
            else:
                formatted_lines.append('</ul>')
        
        return '\n'.join(formatted_lines)

    def _get_html_css(self) -> str:
        """Get CSS styles for HTML report with minimalist design"""
        return """
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }

            body {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Helvetica Neue', Arial, sans-serif;
                line-height: 1.6;
                color: #2c3e50;
                max-width: 210mm;
                margin: 0 auto;
                padding: 20px;
                background: #ffffff;
            }

            .header {
                display: flex;
                justify-content: space-between;
                align-items: center;
                border-bottom: 2px solid #e1e8ed;
                padding-bottom: 15px;
                margin-bottom: 30px;
            }

            .company-logo h1 {
                color: #34495e;
                font-size: 22px;
                font-weight: 400;
                letter-spacing: 0.5px;
            }

            .confidential {
                text-align: right;
                font-size: 11px;
                color: #95a5a6;
                letter-spacing: 0.3px;
            }

            .report-title {
                text-align: center;
                margin-bottom: 35px;
            }

            .report-title h1 {
                font-size: 26px;
                color: #2c3e50;
                margin-bottom: 10px;
                font-weight: 300;
                letter-spacing: 1px;
            }

            .report-info {
                display: flex;
                justify-content: center;
                gap: 40px;
                font-size: 13px;
                color: #7f8c8d;
            }

            .section {
                margin-bottom: 20px;
                border: 1px solid #e1e8ed;
                border-radius: 3px;
                overflow: hidden;
                box-shadow: 0 1px 2px rgba(0,0,0,0.04);
            }

            .section h2 {
                background: #fafbfc;
                color: #2c3e50;
                padding: 14px 20px;
                margin: 0;
                border-bottom: 1px solid #e1e8ed;
                font-size: 16px;
                font-weight: 500;
                letter-spacing: 0.3px;
            }

            /* User Input Section - Minimalist Light Gray */
            .user-input-section {
                border: 1px solid #e8eaec;
                background: #ffffff;
            }

            .user-input-header {
                background: #f8f9fa !important;
                color: #6c757d !important;
                padding: 14px 20px;
                margin: 0;
                border-bottom: 1px solid #e8eaec;
                font-size: 16px;
                font-weight: 500;
                letter-spacing: 0.3px;
            }

            .user-input-box {
                padding: 20px;
                background: #fafbfc;
                margin: 0;
                /* Removed the thick border-left */
            }

            .full-description-container {
                width: 100%;
                overflow: visible;
            }

            .description-label {
                color: #95a5a6;
                font-weight: 500;
                margin-bottom: 12px;
                font-size: 13px;
                text-transform: uppercase;
                letter-spacing: 0.5px;
            }

            .full-description {
                color: #5a6c7d;
                font-size: 14px;
                line-height: 1.8;
                white-space: pre-wrap;
                word-wrap: break-word;
                overflow-wrap: break-word;
                max-width: 100%;
                padding: 16px;
                background: #ffffff;
                border-radius: 2px;
                border: 1px solid #e8eaec;
                min-height: 80px;
                overflow: visible;
            }

            /* AI-Generated Container - Minimalist */
            .ai-generated-container {
                margin-top: 25px;
                padding: 0;
                background: transparent;
                border-radius: 0;
                border: none;
            }

            .ai-main-header {
                color: #34495e;
                font-size: 18px;
                font-weight: 500;
                margin-bottom: 20px;
                padding-bottom: 10px;
                border-bottom: 1px solid #e1e8ed;
                letter-spacing: 0.3px;
            }

            /* AI-Generated Sections - Subtle Blue */
            .ai-section {
                margin-bottom: 20px;
                border: 1px solid #d4dce4;
                border-radius: 3px;
                overflow: hidden;
                background: #ffffff;
                box-shadow: 0 1px 2px rgba(0,0,0,0.04);
            }

            .ai-header {
                background: #516179 !important;
                color: white !important;
                padding: 14px 20px;
                margin: 0;
                font-size: 16px;
                font-weight: 500;
                border-bottom: none;
                letter-spacing: 0.3px;
            }

            .ai-content {
                padding: 20px;
                background: #ffffff;
            }

            .ai-content p {
                margin: 12px 0;
                line-height: 1.7;
                color: #4a5568;
                font-size: 14px;
            }

            .ai-content ol, .ai-content ul {
                margin: 15px 0 15px 20px;
            }

            .ai-content li {
                margin-bottom: 8px;
                line-height: 1.6;
                color: #4a5568;
                font-size: 14px;
            }

            .info-table {
                width: 100%;
                border-collapse: collapse;
                margin: 0;
            }

            .info-table td {
                padding: 12px 15px;
                border-bottom: 1px solid #f0f3f5;
                vertical-align: top;
                font-size: 14px;
            }

            .info-table td.label {
                background: #fafbfc;
                font-weight: 500;
                width: 25%;
                color: #6c757d;
                font-size: 13px;
                text-transform: uppercase;
                letter-spacing: 0.3px;
            }

            .priority-critical {
                background: #e74c3c;
                color: white;
                padding: 4px 10px;
                border-radius: 2px;
                font-weight: 500;
                font-size: 11px;
                text-transform: uppercase;
                letter-spacing: 0.5px;
            }

            .priority-major {
                background: #e67e22;
                color: white;
                padding: 4px 10px;
                border-radius: 2px;
                font-weight: 500;
                font-size: 11px;
                text-transform: uppercase;
                letter-spacing: 0.5px;
            }

            .priority-minor {
                background: #27ae60;
                color: white;
                padding: 4px 10px;
                border-radius: 2px;
                font-weight: 500;
                font-size: 11px;
                text-transform: uppercase;
                letter-spacing: 0.5px;
            }

            .footer {
                margin-top: 40px;
                padding-top: 20px;
                border-top: 1px solid #e1e8ed;
                text-align: center;
                font-size: 11px;
                color: #95a5a6;
                letter-spacing: 0.3px;
            }

            .footer p {
                margin: 5px 0;
            }

            @media print {
                body {
                    max-width: none;
                    margin: 0;
                    padding: 15px;
                }

                .ai-section {
                    page-break-inside: avoid;
                    break-inside: avoid;
                }

                .section {
                    page-break-inside: avoid;
                    break-inside: avoid;
                }
                
                .full-description {
                    page-break-inside: auto;
                }
            }
        """

    def _get_pdf_css(self) -> str:
        """Get CSS styles optimized for PDF generation"""
        return self._get_html_css() + """
            @page {
                size: A4;
                margin: 2cm;
            }

            body {
                font-size: 11pt;
            }

            .full-description {
                page-break-inside: auto;
                font-size: 10pt;
            }

            .ai-section, .section {
                page-break-inside: avoid;
                break-inside: avoid;
            }

            .ai-content {
                page-break-inside: auto;
            }

            h1, h2, h3 {
                page-break-after: avoid;
                break-after: avoid;
            }

            p, li {
                orphans: 3;
                widows: 3;
            }
        """

    def _add_word_content(self, doc, report: DeviationReport):
        """Add content to Word document with improved formatting"""
        from docx.shared import Inches, RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Add header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"{self.company_name} - CONFIDENTIAL"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title
        title = doc.add_heading('Deviation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Report Information
        doc.add_heading('Report Information', level=1)
        table = doc.add_table(rows=6, cols=4)
        table.style = 'Table Grid'

        # Populate table with information
        info_data = [
            ['Title', report.deviation_info.title, 'Priority', report.deviation_info.priority.value],
            ['Status', report.deviation_info.status.value, 'Type', report.deviation_info.deviation_type.value if report.deviation_info.deviation_type else 'Not specified'],
            ['Date', report.deviation_info.date_of_occurrence.strftime("%m/%d/%Y") if report.deviation_info.date_of_occurrence else 'Not specified', 'Department', ", ".join([d.value for d in report.deviation_info.department])],
            ['Batch/Lot', report.deviation_info.batch_lot_number or 'Not applicable', 'Quantity', report.deviation_info.quantity_impacted or 'Not specified'],
            ['Planned', 'Yes' if report.deviation_info.is_planned_deviation else 'No', 'Initiator', report.header.initiator or 'Not specified'],
            ['Record ID', report.header.record_id, 'Generated', report.header.report_generated.strftime("%m/%d/%Y %I:%M %p")]
        ]

        for i, (col1, val1, col2, val2) in enumerate(info_data):
            cells = table.rows[i].cells
            cells[0].text = col1
            cells[1].text = val1
            cells[2].text = col2
            cells[3].text = val2

        # User Input Section with gray background
        user_heading = doc.add_heading('👤 User Input - Original Deviation Description', level=1)
        # Make heading gray
        for run in user_heading.runs:
            run.font.color.rgb = RGBColor(108, 117, 125)  # Gray color
        
        # Add full description with gray text
        user_para = doc.add_paragraph()
        user_run = user_para.add_run(report.deviation_info.description or 'No description provided.')
        user_run.font.color.rgb = RGBColor(73, 80, 87)  # Darker gray for text
        user_run.font.size = Pt(11)
        
        # Add spacing
        doc.add_paragraph()

        # AI-Generated Sections with blue headers
        ai_heading = doc.add_heading('🤖 AI-Generated Analysis', level=1)
        for run in ai_heading.runs:
            run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
        
        ai_content = getattr(report.ai_generated, 'enhanced_description', '')
        ai_sections = self._extract_ai_sections(ai_content)

        section_order = ["Deviation Summary", "Event Timeline", "Root Cause Analysis", "Impact Assessment", "CAPA Plan"]
        section_icons = {
            "Deviation Summary": "📋",
            "Event Timeline": "⏰", 
            "Root Cause Analysis": "🔍",
            "Impact Assessment": "⚠️",
            "CAPA Plan": "🔧"
        }

        for section_name in section_order:
            if section_name in ai_sections and ai_sections[section_name]:
                icon = section_icons.get(section_name, "🤖")
                
                # Add section heading with blue color
                section_heading = doc.add_heading(f'{icon} {section_name}', level=2)
                for run in section_heading.runs:
                    run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
                
                # Add section content
                doc.add_paragraph(ai_sections[section_name])

        # Add footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Report Generated: {datetime.now().strftime('%m/%d/%Y %I:%M %p')} | {self.company_name} - CONFIDENTIAL"